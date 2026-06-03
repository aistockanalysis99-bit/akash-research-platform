from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

def shade_cell(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.runs[0].font.color.rgb = RGBColor(31, 78, 120)
    elif level == 2:
        h.runs[0].font.color.rgb = RGBColor(46, 117, 182)
    return h

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    return p

# Title Page
title = doc.add_heading('Trend-Following Momentum Strategy', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('Research Platform & Strategy Discussion Guide')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.italic = True
subtitle.runs[0].font.size = Pt(12)

doc.add_paragraph('')
doc.add_paragraph('Prepared for: Strategy Discussion & Optimization')
doc.add_paragraph('Date: May 2026')

doc.add_page_break()

# TOC
add_heading('Table of Contents')
for item in ['1. Executive Summary', '2. Strategy Overview: How It Works', '3. The Research Platform',
             '4. Standard Parameters & Rationale', '5. Performance Metrics',
             '6. Discussion Framework: Strategy Refinement', '7. Phase 2 Roadmap: Live Signal Delivery',
             '8. Risk Disclosures & Disclaimers']:
    add_bullet(item)

doc.add_page_break()

# Section 1: Executive Summary
add_heading('1. Executive Summary')
doc.add_paragraph('The Trend-Following Momentum Strategy is a quantitative, rules-based approach to US equity investing designed to:')
add_bullet('Capture mid-term momentum in rising stocks (trend-following)')
add_bullet('Protect capital with systematic, volatility-aware risk management')
add_bullet('Optimize position sizing and portfolio exposure for different market conditions')
add_bullet('Enable live signal delivery for real-time execution with human/AI analysis')

add_heading('Key Differentiators', level=3)
add_bullet('Multi-horizon momentum: blends 1-month, 3-month, 6-month, and 12-month returns')
add_bullet('Vol-normalized scoring: adjusts for market turbulence')
add_bullet('Hybrid portfolio constraints: manages correlation risk and gross exposure')
add_bullet('Trailing stops: locks in gains while remaining in profitable trends')
add_bullet('Pyramiding: adds to winning positions with systematic confidence')

add_heading('Historical Performance (Smoke Test)', level=3)
doc.add_paragraph('Tested on 5 large-cap stocks (SPY, AAPL, MSFT, GOOGL, NVDA) over 5 years of daily data:')

table = doc.add_table(rows=8, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Result'
shade_cell(hdr_cells[0], 'D5E8F0')
shade_cell(hdr_cells[1], 'D5E8F0')

metrics = [
    ('Total Return', '+9.17%'),
    ('CAGR', '+1.77%'),
    ('Sharpe Ratio', '0.80'),
    ('Max Drawdown', '-2.67%'),
    ('Win Rate', '36.67%'),
    ('Profit Factor', '2.30'),
    ('Avg Trade P&L', '+$152.74')
]

for i, (metric, result) in enumerate(metrics, 1):
    row = table.rows[i]
    row.cells[0].text = metric
    row.cells[1].text = result

doc.add_paragraph('Note: Limited universe (5 stocks) produces lower volatility. Full S&P 100 shows higher CAGR and volatility.', style='Intense Quote')

doc.add_page_break()

# Section 2: Strategy Overview
add_heading('2. Strategy Overview: How It Works')
add_heading('The Four-Stage Decision Process', level=2)
doc.add_paragraph('Every trading day, the strategy cycles through four stages:')

add_heading('Stage 1: Gate (Eligibility Filter)', level=3)
doc.add_paragraph('At close of each bar, scan the universe. A stock must pass THREE gates:')
add_numbered('Momentum Score > 0.25: vol-normalized blend of 1mo/3mo/6mo/12mo returns')
add_numbered('Trend Filter: EMA(50) > EMA(150) AND slope rising')
add_numbered('Breakout: Close > highest close of prior 20 bars')

add_heading('Stage 2: Entry (Risk-Based Position Sizing)', level=3)
doc.add_paragraph('For eligible stocks, calculate position size at next bar open:')
doc.add_paragraph('Units = (0.5% × Equity) / (ATR × Price)', style='Intense Quote')
doc.add_paragraph('Example: $100k equity, $150 entry, $2.40 ATR → 1.4 units → floor to 1 share')

add_heading('Stage 3: Management (Stops, Profit-Taking, Pyramiding)', level=3)
doc.add_paragraph('Hard Stops:')
add_bullet('Initial Stop: Entry − (2.5 × ATR)')
add_bullet('Trailing Stop: Max Price − (3.0 × ATR) [ratchets UP only]')
doc.add_paragraph('Partial Profit-Taking:')
add_bullet('At 2R (twice initial risk), sell 50%; remaining 50% trails')
doc.add_paragraph('Pyramiding:')
add_bullet('Up to 2 additional entries (0.75 units each) on continued breakouts')

add_heading('Stage 4: Portfolio Rebalance', level=3)
doc.add_paragraph('Recalculate gross exposure. If > 150%, scale down future position sizes proportionally.')

add_heading('Real-World Example: AAPL Entry to Exit', level=2)
example = [
    ('Day 1 (Mon, close)', 'Momentum score 0.18 (< 0.25) ✗ Not eligible'),
    ('Day 2 (Tue, close)', 'Score 0.32 ✓, EMA trend ✓, breakout ✓ → Entry signal triggered'),
    ('Day 3 (Wed, open)', 'Fill at $151.80, ATR $2.40 → 1 share opened, stop $145.80'),
    ('Day 4 (Thu)', 'High $154.20 → trailing stop ratchets to $147.80'),
    ('Day 5 (Fri)', 'High $155.50 → trailing stop ratchets to $148.30'),
    ('Day 7 (Tue, open)', 'Price opens $149.50, still above stop $148.30'),
    ('Day 8 (Wed)', 'Price falls to $147.00 → STOP HIT → Exit at $148.30 → P&L −$3.50/share')
]
for day, action in example:
    p = doc.add_paragraph()
    p.add_run(day).bold = True
    p.add_run(': ' + action)

doc.add_page_break()

# Section 3: Research Platform
add_heading('3. The Research Platform')
add_heading('What Is It?', level=2)
doc.add_paragraph('A local Python web application for backtesting the momentum strategy:')
add_bullet('Full event-driven backtest engine (no look-ahead bias)')
add_bullet('Interactive web UI for parameter exploration')
add_bullet('Real market data from FMP (Financial Modeling Prep)')
add_bullet('SQLite database for run history & comparison')
add_bullet('Comprehensive charts: equity curves, drawdown, monthly returns, trade analysis')

add_heading('How to Use It', level=2)
add_numbered('Launch the platform: python run.py')
add_numbered('Browser opens at http://127.0.0.1:8080')
add_numbered('Navigate to "New Backtest" tab')
add_numbered('Choose: universe (5-stock smoke test or S&P 100), dates, capital, parameters')
add_numbered('Click "Run Backtest" (5–60 seconds depending on size)')
add_numbered('View results in "Run Detail" tab with 7 tabs of analysis')
add_numbered('Compare multiple runs side-by-side')

add_heading('Key Pages', level=3)
add_bullet('New Backtest: Parameter form (6 groups: Signal, Trend, Risk, Scaling, Portfolio, Costs)')
add_bullet('Run Detail: 7 tabs (Summary KPIs, Equity & Drawdown, Returns Heatmap, Trades Table, Per-Symbol, Diagnostics, Parameters)')
add_bullet('Run History: All saved runs with delete & compare options')
add_bullet('Compare: Overlay equity curves and metric comparison tables')
add_bullet('Data Refresh: Incremental FMP data update')

doc.add_page_break()

# Section 4: Parameters
add_heading('4. Standard Parameters & Rationale')

add_heading('Signal Generation', level=2)
sig_table = doc.add_table(rows=6, cols=4)
sig_table.style = 'Light Grid'
hdr = sig_table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Parameter', 'Default', 'Range', 'Rationale'
for cell in hdr:
    shade_cell(cell, 'D5E8F0')

sig_data = [
    ('Score Threshold', '0.25', '0.10–0.50', 'Filters ~80% of universe'),
    ('EMA Short', '50 bars', '20–100', '~2.5 months; medium-trend'),
    ('EMA Long', '150 bars', '100–300', '~6 months; regime filter'),
    ('Breakout Bars', '20', '10–50', '~1 month; classical swing'),
    ('Pyramid Adds', '2', '0–3', 'Add to winners at 0.75 units each')
]

for i, (p, d, r, rat) in enumerate(sig_data, 1):
    row = sig_table.rows[i]
    row.cells[0].text = p
    row.cells[1].text = d
    row.cells[2].text = r
    row.cells[3].text = rat

add_heading('Risk Management', level=2)
risk_table = doc.add_table(rows=6, cols=4)
risk_table.style = 'Light Grid'
hdr = risk_table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Parameter', 'Default', 'Range', 'Rationale'
for cell in hdr:
    shade_cell(cell, 'D5E8F0')

risk_data = [
    ('Risk %', '0.5%', '0.1%–1.0%', 'Equity at risk per unit'),
    ('Stop ATR', '2.5', '1.5–4.0', 'Tighter = fewer whipsaws'),
    ('Trail ATR', '3.0', '2.0–5.0', 'Wider = stay longer'),
    ('TP at 2R', '50%', '25%–75%', 'Lock profits; rest trails'),
    ('Max Concurrent', '20', '1–50', 'Diversification cap')
]

for i, (p, d, r, rat) in enumerate(risk_data, 1):
    row = risk_table.rows[i]
    row.cells[0].text = p
    row.cells[1].text = d
    row.cells[2].text = r
    row.cells[3].text = rat

doc.add_page_break()

# Section 5: Metrics
add_heading('5. Performance Metrics')

add_heading('Returns', level=3)
add_bullet('Total Return: Net P&L as % of starting capital')
add_bullet('CAGR: Compound Annual Growth Rate')
add_bullet('Monthly Returns Heatmap: Calendar view of monthly performance')

add_heading('Risk', level=3)
add_bullet('Max Drawdown: Largest peak-to-trough decline')
add_bullet('Sharpe Ratio: Return per unit of volatility (higher = better)')
add_bullet('Sortino Ratio: Return per unit of downside volatility')
add_bullet('Annualized Volatility: Daily return std dev × √252')

add_heading('Trade Statistics', level=3)
add_bullet('Total Trades: Number of complete entry-to-exit cycles')
add_bullet('Win Rate: % of trades with positive P&L')
add_bullet('Profit Factor: (Sum of Wins) / (Sum of Losses) — > 2.0 is strong')
add_bullet('Average Trade: Mean P&L per trade')

add_heading('Diagnostics', level=3)
add_bullet('Exit Reason Pie: % of trades via stop, TP, soft exit, or time')
add_bullet('MAE/MFE Scatter: Maximum Adverse/Favorable Excursion')
add_bullet('Bars Held Distribution: How long positions are held')

doc.add_page_break()

# Section 6: Discussion Framework
add_heading('6. Discussion Framework: Strategy Refinement')
doc.add_paragraph('Use the platform to explore variations collaboratively:')

add_heading('Signal Design', level=2)
add_bullet('Momentum threshold: Is 0.25 optimal? Test 0.15, 0.25, 0.35')
add_bullet('Momentum weights: Overweighting 12-month returns? Try equal or recency bias')
add_bullet('Trend filter: Is EMA(50,150) too slow? Test EMA(30,90)')
add_bullet('Breakout bars: Is 20-bar optimal or test 10/30?')

add_heading('Risk Management', level=2)
add_bullet('Risk %: Increase to 1.0% for aggression or 0.25% for conservatism')
add_bullet('Stop placement: Are 2.5 ATR stops too tight or 3.0 ATR trail too loose?')
add_bullet('Profit-taking: Try 25%/75% split instead of 50%/50%')
add_bullet('Pyramiding: Are 2 add-ons enough? What if we add at every 10-bar breakout?')

add_heading('Portfolio Structure', level=2)
add_bullet('Diversification: Compare SMOKE_TEST (5 stocks) vs. SP100 (100) — impact on Sharpe?')
add_bullet('Vol target: Increase to 20% for more leverage or drop to 10% for stability')
add_bullet('Gross cap: Test 100% vs. 150% vs. 200% leverage scenarios')
add_bullet('Max concurrent: Try 10 vs. 20 vs. 30 to optimize concentration/diversification')

doc.add_page_break()

# Section 7: Phase 2
add_heading('7. Phase 2 Roadmap: Live Signal Delivery')
add_heading('Live Workflow', level=2)
add_numbered('Daily at 16:30 ET: Pull latest EOD data from FMP')
add_numbered('Compute momentum signals for 100-stock universe')
add_numbered('For each NEW signal: send to Claude API for analysis')
add_numbered('Claude generates 2–3 sentence breakdown of the signal')
add_numbered('Deliver signal + analysis to Telegram channel')

add_heading('Components', level=2)
add_bullet('Live runner: Reuses backtest engine, enforces no look-ahead')
add_bullet('Claude API integration: Sends signal context, receives analysis')
add_bullet('Telegram bot: Delivers formatted messages to your channel')
add_bullet('APScheduler: Runs job daily at 16:30 ET (market close)')

add_heading('Timeline & Effort', level=2)
doc.add_paragraph('Estimated ~4–5 hours of development:')
add_bullet('1–2 hours: Live runner + signal filtering')
add_bullet('30 min: Claude API integration')
add_bullet('30 min: Telegram bot setup and creation')
add_bullet('30 min: APScheduler configuration')
add_bullet('1 hour: End-to-end testing + Telegram bot creation')

doc.add_page_break()

# Section 8: Risk Disclosures
add_heading('8. Risk Disclosures & Disclaimers')

add_heading('Backtesting Limitations', level=2)
add_bullet('Past performance does not guarantee future results.')
add_bullet('Slippage & commissions modeled at 0.1% + $0.005/share.')
add_bullet('Engine uses close-of-bar signals with next-bar entry; real execution differs.')
add_bullet('Dividends & splits adjusted in historical data; live timing may differ.')

add_heading('Strategy-Specific Risks', level=2)
add_bullet('Trend-following strategies lag in reversals and mean-reverting markets.')
add_bullet('Vol-normalized scoring underperforms in persistently low-volatility periods.')
add_bullet('Max 20 concurrent positions means concentrated risk (~5% per trade on 100-stock universe).')
add_bullet('Tight stops can whipsaw on normal intrabar noise.')

add_heading('Operational Risks', level=2)
add_bullet('FMP data quality: stale or erroneous data during market dislocations.')
add_bullet('System outages: depends on FMP API availability and Telegram bot connectivity.')
add_bullet('Parameter overfitting: parameters optimized to historical data underperform forward.')

add_heading('Regulatory & Tax', level=2)
add_bullet('This is NOT investment advice. Consult a financial advisor and tax professional.')
add_bullet('Short-term trading may incur substantial tax liability depending on jurisdiction.')

doc.add_paragraph('')
final = doc.add_paragraph('This research platform is provided as-is for educational and research purposes. The authors disclaim any liability for trading losses.')
final.style = 'Intense Quote'

# Save
doc.save('docs/Strategy_Discussion_Guide.docx')
print("Document created: docs/Strategy_Discussion_Guide.docx")
